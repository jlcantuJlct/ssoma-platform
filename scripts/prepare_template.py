import os
import re
import unicodedata
from docx import Document
from docx.shared import Pt

def normalize_tag(text):
    text = text.lower()
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('utf-8')
    text = re.sub(r'[^a-z0-9]', '_', text)
    text = re.sub(r'_+', '_', text)
    return text.strip('_')

def main():
    input_path = r"C:\Users\jlcan\Desktop\CASA 2026\Informe San Clemente\Marzo 2026\PAD_SAN CLEMENTE - 12.04.26 B VF2.docx"
    output_path = r"C:\Users\jlcan\Desktop\Seguimiento de plataforma de seguridad Antigravity\ssoma-platform\public\templates\Plantilla_Base.docx"

    print("Cargando documento original...")
    doc = Document(input_path)

    print("Reemplazando fechas y meses...")
    # Reemplazar Marzo -> {MES_REPORTE}
    for p in doc.paragraphs:
        for run in p.runs:
            if 'Marzo' in run.text:
                run.text = run.text.replace('Marzo', '{MES_REPORTE}')
            if 'marzo' in run.text:
                run.text = run.text.replace('marzo', '{MES_REPORTE}')
            if 'MARZO' in run.text:
                run.text = run.text.replace('MARZO', '{MES_REPORTE}')
            if '2026' in run.text:
                run.text = run.text.replace('2026', '{ANIO_REPORTE}')

    print("Eliminando imágenes antiguas e inyectando etiquetas de automatización...")
    NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Iterate all paragraphs to process images and insert tags
    for i, p in enumerate(doc.paragraphs):
        # 1. Eliminar imágenes de los runs de este párrafo
        for run in p.runs:
            drawings = run._element.findall('.//w:drawing', namespaces=NSMAP)
            for d in drawings:
                parent = d.getparent()
                if parent is not None:
                    parent.remove(d)

        # 2. Add docxtemplater tags if paragraph contains "Fotografía...:"
        text = p.text
        if 'fotografía' in text.lower() and ':' in text:
            # Extract caption
            parts = text.split(':')
            if len(parts) > 1:
                description = parts[1].strip()
                tag = normalize_tag(description.replace('.', ''))
                
                new_p = p.insert_paragraph_before()
                p._p.addnext(new_p._p)
                
                new_p.text = f"{{#{tag}}}\n{{%url}}\n{{/{tag}}}"
                new_p.style = p.style
                for r in new_p.runs:
                    r.font.size = Pt(10)

    # También revisar tablas por si hay imágenes incrustadas dentro de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        drawings = run._element.findall('.//w:drawing', namespaces=NSMAP)
                        for d in drawings:
                            parent = d.getparent()
                            if parent is not None:
                                parent.remove(d)
                            
                    # Si alguna tabla tiene etiquetas de foto
                    text = p.text
                    if 'fotografía' in text.lower() and ':' in text:
                        parts = text.split(':')
                        if len(parts) > 1:
                            description = parts[1].strip()
                            tag = normalize_tag(description.replace('.', ''))
                            p.add_run(f"\n{{#{tag}}}\n{{%url}}\n{{/{tag}}}")

    # Guardar documento local en public/templates
    print(f"Guardando plantilla maestra en {output_path}...")
    doc.save(output_path)
    print("¡Plantilla procesada exitosamente!")

if __name__ == '__main__':
    main()
