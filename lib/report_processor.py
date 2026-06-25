import sys
import json
import os
import re
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

def normalize_text(text):
    if not text: return ""
    # Quitar tildes y convertir a minúsculas
    text = text.lower()
    text = re.sub(r'[áéíóú]', lambda m: {'á':'a','é':'e','í':'i','ó':'o','ú':'u'}[m.group()], text)
    return re.sub(r'[^a-z0-9]', ' ', text).strip()

def process_report(template_path, output_path, data_file_path):
    print(f"🤖 Iniciando Procesador Inteligente...")
    
    # Leer datos desde el archivo JSON
    with open(data_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    month_name = data.get("monthName", "Mes").upper()
    year = str(data.get("year", "2026"))
    evidence = data.get("evidence", [])
    
    print(f"📅 Periodo: {month_name} {year}")
    print(f"📸 Evidencias disponibles: {len(evidence)}")

    doc = Document(template_path)
    
    # 1. REEMPLAZO DINÁMICO DE FECHAS
    # Buscamos meses del 2025 o 2026 para actualizar al actual
    meses_busqueda = ["ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO", "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"]
    
    for p in doc.paragraphs:
        original_text = p.text
        for m in meses_busqueda:
            if m in p.text.upper():
                # Reemplazar mes y año
                p.text = p.text.upper().replace(m, month_name).replace("2025", year).replace("2026", year)
                # Mantener capitalización original si es posible (simple check)
                if original_text.istitle(): p.text = p.text.title()
                elif original_text.islower(): p.text = p.text.lower()

    # 2. INYECCIÓN INTELIGENTE DE IMÁGENES
    used_photos = set()

    for p in doc.paragraphs:
        p_text = p.text.strip()
        p_norm = normalize_text(p_text)
        
        # Criterios para identificar dónde va una foto:
        # - El párrafo contiene "FOTO", "FOTOGRAFÍA", "GRÁFICO"
        # - O el párrafo es un "Código" manual como [[FOTO:Nombre]]
        is_placeholder = any(k in p_text.upper() for k in ["FOTO", "FOTOGRAFÍA", "GRÁFICO"])
        manual_code = re.search(r'\[\[FOTO:(.*?)\]\]', p_text)

        if is_placeholder or manual_code:
            search_query = manual_code.group(1) if manual_code else p_text
            search_norm = normalize_text(search_query)
            
            print(f"🔍 Buscando coincidencia para: '{search_query[:40]}...'")
            
            best_match = None
            for photo in evidence:
                if photo.get("file_url") in used_photos: continue
                
                # Campos para comparar
                fields = [
                    photo.get("activity", ""),
                    photo.get("description", ""),
                    photo.get("category", ""),
                    photo.get("zona", ""),
                    photo.get("location", "")
                ]
                
                photo_content = normalize_text(" ".join(filter(None, fields)))
                
                # Si las palabras clave del Word están en la descripción de la foto
                # O viceversa
                keywords = [w for w in search_norm.split() if len(w) > 3]
                if any(k in photo_content for k in keywords):
                    best_match = photo
                    break
            
            if best_match:
                img_url = best_match.get("file_url")
                print(f"   ✅ MATCH: '{best_match.get('activity')}'")
                
                try:
                    # Convertir a link directo de Drive si es necesario
                    if "drive.google.com" in img_url:
                        file_id = img_url.split('id=')[-1].split('&')[0] if 'id=' in img_url else img_url.split('/')[-2]
                        img_url = f"https://docs.google.com/uc?export=download&id={file_id}"
                    
                    response = requests.get(img_url, timeout=15)
                    if response.status_code == 200:
                        img_stream = BytesIO(response.content)
                        # Insertar imagen ANTES del pie de foto
                        new_p = p.insert_paragraph_before('')
                        run = new_p.add_run()
                        # Tamaño uniforme: 7 cm de ancho × 4 cm de alto (igual que el Generador Dinámico)
                        from docx.shared import Cm
                        run.add_picture(img_stream, width=Cm(7), height=Cm(4))
                        used_photos.add(best_match.get("file_url"))
                        
                        # Si era un código manual [[...]], borrar el código
                        if manual_code:
                            p.text = p.text.replace(manual_code.group(0), "").strip()
                except Exception as e:
                    print(f"   ❌ ERROR al insertar: {e}")

    # Guardar resultado
    doc.save(output_path)
    print(f"🏁 Informe generado exitosamente: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Uso: python report_processor.py <template> <output> <data_json>")
    else:
        process_report(sys.argv[1], sys.argv[2], sys.argv[3])
